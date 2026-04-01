import asyncio
import json
import math
import os
import re
from collections import Counter
from concurrent.futures import ProcessPoolExecutor

try:
    import jieba
except Exception:  # pragma: no cover - fallback for minimal environments
    jieba = None

try:
    from rank_bm25 import BM25Okapi
except Exception:  # pragma: no cover - fallback for minimal environments
    BM25Okapi = None


LEGAL_CLAUSE_PATTERN = re.compile(r'第[一二三四五六七八九十百千零\d]+[条章节款编]')
SENTENCE_SPLIT_PATTERN = re.compile(r'[。；！？\n]')

CHINESE_DATE_PATTERN = re.compile(
    r'(?:(?:19|20)\d{2}[年\-/]\d{1,2}[月\-/]\d{1,2}[日]?)|'
    r'(?:\d{1,2}[月]\d{1,2}[日]?)|'
    r'(?:(?:19|20)\d{2}年)|'
    r'(?:\d{4}[-/]\d{1,2}[-/]\d{1,2})'
)
CURRENCY_AMOUNT_PATTERN = re.compile(
    r'(?:(?:人民币|美元|欧元|英镑|港币|日元|澳元)?\s*(?:[\d,，.]+)\s*(?:元|万|亿|美元|欧元|英镑|港币|日元|澳元|万(?:美元|欧元|英镑|港币|日元|澳元)))|'
    r'(?:(?:租金|押金|违约金|赔偿|费用|总额|金额|付款)?[:：]?\s*[\d,，.]+\s*(?:元|万|亿|美元|欧元))'
)
PARTY_PATTERN = re.compile(
    r'(?:甲方|乙方|出租人|承租人|房东|Tenant| Landlord|当事人|委托人|对方(?:当事人)?)[:：]?\s*([^\s，。、；。]{2,30})'
)
EMAIL_PATTERN = re.compile(r'[\w.+-]+@[\w-]+\.[\w.-]+')
PHONE_PATTERN = re.compile(r'(?:1[3-9]\d{9}|(?:\+?86)?[-\s()]?\d{3,4}[-\s()]?\d{3,4}[-\s()]?\d{0,4})')
ADDRESS_PATTERN = re.compile(r'(?:地址|住所|Location)[:：]?\s*([^，。；\n]{5,60})')


def _segment(text):
    if jieba is None:
        tokens = re.findall(r"[\w\u4e00-\u9fff]+", text.lower())
        return [t for t in tokens if t.strip()]
    tokens = list(jieba.cut_for_search(text))
    return [t for t in tokens if t.strip()]


def _extract_text_pymupdf(file_path):
    try:
        import pymupdf
        doc = pymupdf.open(file_path)
        pages = []
        for page in doc:
            pages.append(page.get_text())
        doc.close()
        return "\n".join(pages)
    except Exception as exc:
        return None, f"PyMuPDF extraction failed: {exc}"


def _extract_text_pypdf(file_path):
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return text
    except Exception as exc:
        return None, f"pypdf extraction failed: {exc}"


def _read_file(file_path):
    lower_path = file_path.lower()
    if lower_path.endswith((".txt", ".md", ".markdown")):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as handle:
            return handle.read()
    if lower_path.endswith(".pdf"):
        result = _extract_text_pymupdf(file_path)
        if isinstance(result, tuple):
            result = _extract_text_pypdf(file_path)
            if isinstance(result, tuple):
                return result
        return result
    if lower_path.endswith(".docx"):
        try:
            import docx

            document = docx.Document(file_path)
            parts = []
            for p in document.paragraphs:
                if p.text and p.text.strip():
                    parts.append(p.text)
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        parts.append(row_text)
            return "\n".join(parts) if parts else (None, "docx: no readable content")
        except Exception as exc:
            return None, f"docx read failed: {exc}"
    return None, "Unsupported file format"


# Indexed text must not be OCR/read error placeholders (they pollute search).
_OCR_OR_READ_FAILURE_MARKERS = (
    "[GLM-OCR Error]",
    "[GLM-OCR produced no usable",
    "[Fallback scan error]",
    "[Unsupported:",
    "[DOCX read error]",
    "[Vision OCR fallback",
    "[No text extracted",
)


def _is_unindexable_placeholder_text(content: str) -> bool:
    s = (content or "").strip()
    if not s:
        return True
    if s.startswith("GenerationResult("):
        return True
    head = s[:1200]
    for marker in _OCR_OR_READ_FAILURE_MARKERS:
        if marker in head:
            return True
    if s.startswith("[") and len(s) < 800:
        low = s.lower()
        if any(
            w in low
            for w in ("error", "failed", "unsupported", "unavailable", "no module named")
        ):
            return True
    return False


def _validate_content_for_indexing(content: str) -> tuple[bool, str]:
    if _is_unindexable_placeholder_text(content):
        return (
            False,
            "Indexing skipped: no real document text (OCR or file read failed). "
            "Install `pymupdf` and `mlx-vlm` in the venv, then re-upload or use POST "
            "`/documents/{filename}/rescan`.",
        )
    return True, ""


def _structure_aware_chunk(text, chunk_size=512, overlap=80):
    cleaned = re.sub(r"\s+", " ", text).strip()
    if not cleaned:
        return []

    clause_splits = LEGAL_CLAUSE_PATTERN.split(cleaned)
    clause_markers = LEGAL_CLAUSE_PATTERN.findall(cleaned)

    chunks = []
    current_chunk = ""

    if clause_markers:
        segments = []
        for i, segment in enumerate(clause_splits):
            if i == 0 and segment.strip():
                segments.append(segment.strip())
            if i < len(clause_markers):
                segments.append(clause_markers[i] + (clause_splits[i + 1] if i + 1 < len(clause_splits) else ""))
            elif segment.strip():
                segments.append(segment.strip())

        for segment in segments:
            if len(current_chunk) + len(segment) <= chunk_size:
                current_chunk += segment
            else:
                if current_chunk:
                    chunks.append(current_chunk)
                if len(segment) > chunk_size:
                    sub_chunks = _fixed_chunk(segment, chunk_size, overlap)
                    chunks.extend(sub_chunks)
                    current_chunk = ""
                else:
                    current_chunk = segment
    else:
        sentences = [s.strip() for s in SENTENCE_SPLIT_PATTERN.split(cleaned) if s.strip()]
        for sentence in sentences:
            if len(current_chunk) + len(sentence) <= chunk_size:
                current_chunk += sentence + "。"
            else:
                if current_chunk:
                    chunks.append(current_chunk)
                current_chunk = sentence + "。"

    if current_chunk:
        chunks.append(current_chunk)
    return chunks


def _fixed_chunk(text, chunk_size=512, overlap=80):
    cleaned = re.sub(r"\s+", " ", text).strip()
    if not cleaned:
        return []
    chunks = []
    start = 0
    text_length = len(cleaned)
    while start < text_length:
        end = min(start + chunk_size, text_length)
        chunks.append(cleaned[start:end])
        if end >= text_length:
            break
        start = max(0, end - overlap)
    return chunks


def _normalize_lines(text: str) -> list[str]:
    return [line.strip() for line in text.splitlines() if line.strip()]


def _extract_pdf_metadata(file_path):
    import pymupdf
    doc = pymupdf.open(file_path)
    meta = {
        "page_count": doc.page_count,
        "title": doc.metadata.get("title", ""),
        "author": doc.metadata.get("author", ""),
        "subject": doc.metadata.get("subject", ""),
        "creator": doc.metadata.get("creator", ""),
        "producer": doc.metadata.get("producer", ""),
        "creation_date": doc.metadata.get("creationDate", ""),
        "mod_date": doc.metadata.get("modDate", ""),
        "toc": doc.get_toc(),
        "file_size_bytes": os.path.getsize(file_path),
        "pages": [],
        "total_chars": 0,
    }
    for i in range(doc.page_count):
        page = doc[i]
        text = page.get_text()
        meta["total_chars"] += len(text)
        meta["pages"].append({
            "page_number": i + 1,
            "char_count": len(text),
            "has_images": len(page.get_images()) > 0,
            "width": round(page.rect.width, 1),
            "height": round(page.rect.height, 1),
        })
    doc.close()
    return meta


def _extract_pdf_page_texts(file_path):
    import pymupdf
    doc = pymupdf.open(file_path)
    pages = []
    for i in range(doc.page_count):
        pages.append({"page": i + 1, "text": doc[i].get_text()})
    doc.close()
    return pages


def _render_pdf_page(file_path, page_number=1, dpi=150):
    import pymupdf
    doc = pymupdf.open(file_path)
    idx = page_number - 1
    if not (0 <= idx < doc.page_count):
        doc.close()
        return None
    page = doc[idx]
    zoom = dpi / 72.0
    mat = pymupdf.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=mat)
    png_bytes = pix.tobytes("png")
    doc.close()
    return png_bytes


def _find_text_in_pages(file_path, search_text):
    import pymupdf
    if not search_text or not search_text.strip():
        return []
    doc = pymupdf.open(file_path)
    matches = []
    search_lower = search_text.strip().lower()
    for i in range(doc.page_count):
        page_text = doc[i].get_text().lower()
        if search_lower in page_text:
            matches.append(i + 1)
    doc.close()
    return matches


class RAGManager:
    def __init__(self, persist_directory="./data/vector_store", enable_turboquant=True):
        self.persist_directory = persist_directory
        self.store_path = os.path.join(self.persist_directory, "light_rag_store.json")
        self.documents = []
        self.backend_mode = "lightweight_bm25"
        self.init_error = None
        self.vector_db = None
        self._bm25 = None
        self._tokenized_corpus = []
        self._chunk_meta: list[dict] = []
        # Memory optimization: removed _all_chunks and _all_chunk_meta
        # These were duplicates of data already in self.documents[].chunks
        # Now we iterate documents directly in search
        self._embedding_model = None
        self._use_embeddings = False
        self._search_cache: dict[str, list] = {}
        self._turboquant = None
        self._turboquant_enabled = enable_turboquant
        self._load_store()

    @staticmethod
    def _project_root() -> str:
        return os.path.dirname(__file__)

    @classmethod
    def _normalize_path(cls, path: str) -> str:
        if not path:
            return path
        if os.path.isabs(path):
            return os.path.normpath(path)
        # Persisted paths in older stores may be like "./data/docs/x.pdf".
        return os.path.normpath(os.path.join(cls._project_root(), path.lstrip("./")))

    def _load_store(self):
        os.makedirs(self.persist_directory, exist_ok=True)
        if not os.path.exists(self.store_path):
            return
        try:
            with open(self.store_path, "r", encoding="utf-8") as handle:
                payload = json.load(handle)
            self.documents = payload.get("documents", [])
            # Repair relative paths to keep original file loading stable across restarts.
            changed = False
            for doc in self.documents:
                p = doc.get("path")
                if isinstance(p, str) and p and not os.path.isabs(p):
                    doc["path"] = self._normalize_path(p)
                    changed = True
            # Drop records whose original file is missing to prevent “ghost documents”
            # that cannot be verified/viewed.
            before = len(self.documents)
            kept = []
            for doc in self.documents:
                p = doc.get("path")
                if isinstance(p, str) and p and os.path.exists(p):
                    kept.append(doc)
                else:
                    changed = True
            self.documents = kept
            if before != len(self.documents):
                self.init_error = (
                    f"Removed {before - len(self.documents)} stale RAG document record(s) whose original files are missing."
                )
            if changed:
                self._save_store()
            self._rebuild_index()
        except Exception as exc:
            self.init_error = f"Could not load local RAG store: {exc}"
            self.documents = []

    def _save_store(self):
        payload = {"documents": self.documents}
        with open(self.store_path, "w", encoding="utf-8") as handle:
            json.dump(payload, handle, ensure_ascii=False, indent=2)

    def _rebuild_index(self):
        self._tokenized_corpus = []
        self._chunk_meta = []
        for doc in self.documents:
            fn = doc.get("filename", "")
            for chunk in doc.get("chunks", []):
                self._tokenized_corpus.append(_segment(chunk))
                self._chunk_meta.append({"filename": fn, "chunk": chunk})
        if self._tokenized_corpus and BM25Okapi is not None:
            self._bm25 = BM25Okapi(self._tokenized_corpus)
        else:
            self._bm25 = None
        self._search_cache.clear()
        if self._turboquant_enabled:
            self._init_turboquant()

    def _init_turboquant(self):
        try:
            from turboquant import TurboQuantBM25
            tq = TurboQuantBM25(bit_width=3)
            tq.fit(self._tokenized_corpus)
            self._turboquant = tq
            ratio = tq.compression_ratio
            saved = tq.memory_saved_bytes
            self.backend_mode = f"turboquant_bm25_{tq.bit_width}bit"
            print(f"[RAG] TurboQuant enabled: {ratio:.1f}x compression, {saved} bytes saved")
        except Exception as e:
            print(f"[RAG] TurboQuant unavailable, using standard BM25: {e}")
            self._turboquant = None

    def _init_embeddings(self):
        if self._embedding_model is not None:
            return True
        try:
            from FlagEmbedding import FlagModel
            self._embedding_model = FlagModel(
                "BAAI/bge-m3",
                use_fp16=True,
            )
            self._use_embeddings = True
            self.backend_mode = "hybrid_bm25_dense"
            print("[RAG] bge-m3 embedding model loaded (hybrid mode).")
            return True
        except Exception as e:
            print(f"[RAG] Embedding model unavailable, using BM25 only: {e}")
            self._use_embeddings = False
            return False

    def ensure_ready(self):
        return True

    def ingest_file(self, file_path, **kwargs):
        if not self.ensure_ready():
            return False, "RAG unavailable"

        # If OCR text is provided externally (e.g. from GLM-OCR), use it.
        if "force_ocr_text" in kwargs:
            content = kwargs["force_ocr_text"]
        else:
            content = _read_file(file_path)
            # If direct read failed and it's an image, try OCR
            if isinstance(content, tuple) and not kwargs.get("skip_ocr", False):
                lower_fp = file_path.lower()
                if lower_fp.endswith((".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff", ".gif", ".heic", ".pdf")):
                    try:
                        from ocr_engine import ocr_engine
                        import concurrent.futures
                        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
                            future = executor.submit(asyncio.run, ocr_engine.scan_file(file_path, task="Text Recognition"))
                            content = future.result(timeout=300)
                    except Exception as e:
                        return False, f"OCR failed: {e}"

        if isinstance(content, tuple):
            return False, content[1]

        ok, err_msg = _validate_content_for_indexing(content)
        if not ok:
            return False, err_msg

        chunks = _structure_aware_chunk(content)
        if not chunks:
            return False, "No readable text found in file"

        normalized_path = self._normalize_path(str(file_path))
        filename = os.path.basename(normalized_path)
        lower_fp = file_path.lower()
        if lower_fp.endswith(".pdf"):
            file_type, is_pdf = "pdf", True
        elif lower_fp.endswith((".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff", ".gif", ".heic")):
            file_type, is_pdf = "image", False
        elif lower_fp.endswith(".docx"):
            file_type, is_pdf = "docx", False
        else:
            file_type, is_pdf = "txt", False

        doc_record = {
            "filename": filename,
            "path": normalized_path,
            "chunk_count": len(chunks),
            "preview": chunks[0][:220],
            "chunks": chunks,
            "file_type": file_type,
            "file_size_bytes": os.path.getsize(normalized_path),
            "user_preferences": kwargs.get("user_preferences", {}),
            "ai_metadata": kwargs.get("ai_metadata", {}),
            "ingested_at": __import__("datetime").datetime.now(__import__("datetime").timezone.utc).isoformat(),
        }

        if is_pdf:
            try:
                doc_record["pdf_metadata"] = _extract_pdf_metadata(normalized_path)
            except Exception:
                doc_record["pdf_metadata"] = None
            try:
                doc_record["page_texts"] = _extract_pdf_page_texts(normalized_path)
                # Tag chunks with source page numbers
                page_texts = doc_record["page_texts"]
                tagged_chunks = []
                for chunk in chunks:
                    page_num = self._find_chunk_page(chunk, page_texts)
                    tagged_chunks.append({"text": chunk, "page": page_num})
                doc_record["tagged_chunks"] = tagged_chunks
            except Exception:
                doc_record["page_texts"] = []
                doc_record["tagged_chunks"] = [{"text": c, "page": None} for c in chunks]

        self.documents = [doc for doc in self.documents if doc.get("filename") != filename]
        self.documents.append(doc_record)
        self._save_store()
        self._rebuild_index()
        if self._use_embeddings:
            self._embed_chunks(chunks)
        mode = "hybrid" if self._use_embeddings else "BM25"
        return True, f"Ingested {len(chunks)} chunks with {mode} search ({filename})"

    def _embed_chunks(self, new_chunks):
        pass

    @staticmethod
    def _find_chunk_page(chunk_text, page_texts):
        chunk_lower = chunk_text.strip().lower()[:120]
        for pt in page_texts:
            if chunk_lower in pt["text"].lower():
                return pt["page"]
        return None

    def find_text_in_pages(self, filename, search_text):
        doc = self.get_document_record(filename)
        if doc is None:
            return []
        path = doc.get("path", "")
        if not path or not os.path.exists(path):
            return []
        try:
            return _find_text_in_pages(path, search_text)
        except Exception:
            return []

    def inspect_pdf(self, filename):
        doc = self.get_document_record(filename)
        if doc is None:
            raise KeyError(f"Unknown document: {filename}")
        path = doc.get("path", "")
        if not path or not os.path.exists(path):
            raise ValueError(f"File not found: {path}")
        import pymupdf
        pdf_doc = pymupdf.open(path)
        result = {
            "filename": filename,
            "file_type": "pdf",
            "page_count": pdf_doc.page_count,
            "file_size_bytes": os.path.getsize(path),
            "metadata": {
                "title": pdf_doc.metadata.get("title", ""),
                "author": pdf_doc.metadata.get("author", ""),
                "subject": pdf_doc.metadata.get("subject", ""),
                "creator": pdf_doc.metadata.get("creator", ""),
                "creation_date": pdf_doc.metadata.get("creationDate", ""),
            },
            "toc": pdf_doc.get_toc(),
            "pages": [],
            "total_chars": 0,
            "total_words": 0,
        }
        for i in range(pdf_doc.page_count):
            page = pdf_doc[i]
            text = page.get_text()
            word_count = len(text.split())
            has_tables = False
            try:
                has_tables = len(page.find_tables().tables) > 0
            except Exception:
                pass
            result["pages"].append({
                "page_number": i + 1,
                "char_count": len(text),
                "word_count": word_count,
                "has_images": len(page.get_images()) > 0,
                "has_tables": has_tables,
                "dimensions": f"{page.rect.width:.0f}x{page.rect.height:.0f}",
            })
            result["total_chars"] += len(text)
            result["total_words"] += word_count
        pdf_doc.close()
        return result

    def get_document_record(self, filename: str) -> dict | None:
        for doc in self.documents:
            if doc.get("filename") == filename:
                return doc
        return None

    def get_document_text(self, filename: str) -> str:
        doc = self.get_document_record(filename)
        if doc is None:
            raise KeyError(f"Unknown document: {filename}")
        content = _read_file(doc.get("path", ""))
        if isinstance(content, tuple):
            raise ValueError(content[1])
        return content

    def get_inspect_text(self, filename: str) -> str:
        """Text for inspect/grep: file bytes for TXT/PDF/DOCX, else joined RAG chunks (e.g. OCR'd images)."""
        doc = self.get_document_record(filename)
        if doc is None:
            raise KeyError(f"Unknown document: {filename}")
        try:
            return self.get_document_text(filename)
        except ValueError:
            chunks = doc.get("chunks") or []
            return "\n".join(chunks) if chunks else ""

    def inspect_document(self, filename: str, start_line: int = 1, window: int = 40, max_chars: int = 4000) -> dict:
        text = self.get_inspect_text(filename)
        lines = _normalize_lines(text)
        start_line = max(1, start_line)
        window = max(1, window)
        start_idx = min(len(lines), start_line - 1)
        end_idx = min(len(lines), start_idx + window)
        selected = [
            {"line_number": idx + 1, "text": lines[idx]}
            for idx in range(start_idx, end_idx)
        ]
        excerpt = "\n".join(item["text"] for item in selected)[:max_chars]
        doc = self.get_document_record(filename) or {}
        return {
            "filename": filename,
            "path": doc.get("path"),
            "chunk_count": doc.get("chunk_count", 0),
            "total_lines": len(lines),
            "start_line": start_line,
            "window": window,
            "lines": selected,
            "excerpt": excerpt,
        }

    def grep_document(self, filename: str, pattern: str, max_results: int = 20, case_sensitive: bool = False) -> dict:
        text = self.get_inspect_text(filename)
        lines = _normalize_lines(text)
        flags = 0 if case_sensitive else re.IGNORECASE
        compiled = re.compile(pattern, flags)
        matches = []
        for idx, line in enumerate(lines):
            if compiled.search(line):
                matches.append(
                    {
                        "line_number": idx + 1,
                        "text": line,
                    }
                )
            if len(matches) >= max_results:
                break
        return {
            "filename": filename,
            "pattern": pattern,
            "case_sensitive": case_sensitive,
            "count": len(matches),
            "matches": matches,
        }

    def summarize_document(self, filename: str, max_points: int = 5) -> dict:
        try:
            text = self.get_document_text(filename)
        except ValueError:
            text = self.get_inspect_text(filename)
        doc = self.get_document_record(filename) or {}
        lines = _normalize_lines(text)
        sentences = [s.strip() for s in SENTENCE_SPLIT_PATTERN.split(re.sub(r"\s+", " ", text)) if s.strip()]
        token_counts = Counter(_segment(text))
        scored = []
        for sentence in sentences:
            score = sum(token_counts.get(token, 0) for token in _segment(sentence))
            bonus = 5 if LEGAL_CLAUSE_PATTERN.search(sentence) else 0
            scored.append((score + bonus, sentence))
        scored.sort(key=lambda item: item[0], reverse=True)

        summary_points = []
        seen = set()
        heading_pattern = re.compile(r"^(Section\s+\d+|[一二三四五六七八九十]+、|第[\d一二三四五六七八九十百千]+[条章节款编])", re.IGNORECASE)
        for idx, line in enumerate(lines):
            if heading_pattern.search(line):
                bundle = [line]
                if idx + 1 < len(lines):
                    bundle.append(lines[idx + 1])
                if idx + 2 < len(lines) and len(bundle[1]) < 120:
                    bundle.append(lines[idx + 2])
                point = " ".join(bundle)[:220]
                if point not in seen:
                    summary_points.append(point)
                    seen.add(point)
                if len(summary_points) >= max_points:
                    break
        for _, sentence in scored:
            if sentence in seen:
                continue
            seen.add(sentence)
            summary_points.append(sentence[:220])
            if len(summary_points) >= max_points:
                break

        headings = []
        for line in lines:
            if LEGAL_CLAUSE_PATTERN.search(line) or heading_pattern.search(line) or line.startswith(("一、", "二、", "三、", "四、", "五、", "六、")):
                headings.append(line[:120])
            if len(headings) >= max_points:
                break

        return {
            "filename": filename,
            "chunk_count": doc.get("chunk_count", 0),
            "preview": doc.get("preview", ""),
            "summary_points": summary_points,
            "headings": headings,
        }

    def extract_entities(self, filename: str) -> dict:
        """Extract structured legal entities from a document: dates, amounts, parties, clauses, etc."""
        try:
            text = self.get_document_text(filename)
        except ValueError:
            text = self.get_inspect_text(filename)
        doc = self.get_document_record(filename) or {}

        entities = {
            "dates": [],
            "amounts": [],
            "parties": [],
            "clauses": [],
            "emails": [],
            "phones": [],
            "addresses": [],
        }

        for match in CHINESE_DATE_PATTERN.finditer(text):
            date_str = match.group().strip()
            if date_str not in [d["text"] for d in entities["dates"]]:
                entities["dates"].append({"text": date_str, "position": match.start()})

        for match in CURRENCY_AMOUNT_PATTERN.finditer(text):
            amount_str = match.group().strip()
            if amount_str not in [a["text"] for a in entities["amounts"]]:
                entities["amounts"].append({"text": amount_str, "position": match.start()})

        for match in PARTY_PATTERN.finditer(text):
            party_name = match.group(1).strip()
            if party_name and party_name not in [p["text"] for p in entities["parties"]]:
                entities["parties"].append({"text": party_name, "role": match.group(0)[:3], "position": match.start()})

        for match in EMAIL_PATTERN.finditer(text):
            email = match.group()
            if email not in entities["emails"]:
                entities["emails"].append({"text": email, "position": match.start()})

        for match in PHONE_PATTERN.finditer(text):
            phone = match.group().strip()
            cleaned = re.sub(r'[-\s()]+', '', phone)
            if len(cleaned) >= 7 and cleaned not in [p["text"] for p in entities["phones"]]:
                entities["phones"].append({"text": phone, "position": match.start()})

        for match in ADDRESS_PATTERN.finditer(text):
            addr = match.group(1).strip()
            if addr not in entities["addresses"]:
                entities["addresses"].append({"text": addr, "position": match.start()})

        clause_matches = LEGAL_CLAUSE_PATTERN.finditer(text)
        for match in clause_matches:
            clause_text = match.group().strip()
            start = match.start()
            end = min(start + 100, len(text))
            context = text[start:end]
            sentence_end = context.find("。")
            if sentence_end == -1:
                sentence_end = 50
            clause_content = context[:sentence_end + 1]
            if clause_text not in [c["clause"] for c in entities["clauses"]]:
                entities["clauses"].append({
                    "clause": clause_text,
                    "context": clause_content.strip(),
                    "position": start,
                })

        rent_keywords = ["租金", "月租", "年租", "押金", "保证金", "租期"]
        deposit_section = None
        for keyword in rent_keywords:
            if keyword in text:
                idx = text.index(keyword)
                snippet = text[max(0, idx - 20):idx + 60]
                if deposit_section is None or len(snippet) > len(deposit_section.get("snippet", "")):
                    deposit_section = {"keyword": keyword, "snippet": snippet, "position": idx}

        return {
            "filename": filename,
            "file_type": doc.get("file_type", "unknown"),
            "page_count": (doc.get("pdf_metadata") or {}).get("page_count"),
            "total_chars": len(text),
            "entities": entities,
            "deposit_rent_info": deposit_section,
            "extracted_at": __import__("datetime").datetime.now(__import__("datetime").timezone.utc).isoformat(),
        }

    def search(self, query, k=5):
        cache_key = f"{query}:{k}"
        if cache_key in self._search_cache:
            return self._search_cache[cache_key]
        results = []
        if self._bm25 and self._tokenized_corpus:
            tokenized_query = _segment(query)
            bm25_scores = self._bm25.get_scores(tokenized_query)
            scored_indices = sorted(range(len(bm25_scores)), key=lambda i: bm25_scores[i], reverse=True)
            for idx in scored_indices[:k * 2]:
                score = bm25_scores[idx]
                if score > 0 or (len(self._tokenized_corpus) <= 3 and score > -5):
                    meta = self._chunk_meta[idx] if idx < len(self._chunk_meta) else {}
                    results.append((score, meta.get("filename"), meta.get("chunk")))
        if not results:
            query_tokens = _segment(query)
            for document in self.documents:
                for chunk in document.get("chunks", []):
                    score = self._tf_score(query_tokens, chunk)
                    if score > 0:
                        results.append((score, document.get("filename"), chunk))
        results.sort(key=lambda item: item[0], reverse=True)
        top_chunks = [f"[{filename}] {chunk}" for _, filename, chunk in results[:k]]
        formatted = "\n---\n".join(top_chunks)
        self._search_cache[cache_key] = formatted
        if len(self._search_cache) > 256:
            oldest = next(iter(self._search_cache))
            del self._search_cache[oldest]
        return formatted

    def _substring_chunk_hits(self, query: str, k: int, skip_keys: set | None = None) -> list[dict]:
        """Direct substring / token containment for CJK when BM25 misses (e.g. short OCR chunks)."""
        skip_keys = skip_keys or set()
        q = (query or "").strip()
        if len(q) < 2:
            return []
        tokens = [t for t in _segment(q) if len(t.strip()) >= 2]
        out: list[dict] = []
        for doc in self.documents:
            fn = doc.get("filename") or ""
            for ci, chunk in enumerate(doc.get("chunks") or []):
                if not chunk:
                    continue
                key = (fn, ci)
                if key in skip_keys:
                    continue
                hit = q in chunk
                if not hit and tokens:
                    hit = all(t in chunk for t in tokens[:6])
                if hit:
                    out.append({
                        "filename": fn,
                        "chunk": chunk[:2000],
                        "score": 0.05,
                        "match_mode": "substring",
                    })
                    if len(out) >= k:
                        return out
        return out

    def _fuzzy_match(self, query: str, text: str) -> float:
        """Fuzzy substring match with typo tolerance for Chinese and English."""
        q = query.strip().lower()
        t = text.lower()
        if q in t:
            return 1.0
        if len(q) < 2:
            return 0.0
        q_tokens = _segment(q)
        t_tokens = _segment(t)
        if not q_tokens:
            return 0.0
        matches = sum(1 for qt in q_tokens if any(qt in tt or tt in qt for tt in t_tokens))
        return matches / len(q_tokens) * 0.8

    def _extract_names_from_query(self, query: str) -> list[str]:
        """Extract potential person/entity names from search query."""
        names = []
        q = query.strip()
        if len(q) >= 2:
            tokens = _segment(q)
            for t in tokens:
                if len(t) >= 2 and re.match(r'^[\u4e00-\u9fff]+$', t):
                    names.append(t)
        return names

    def _extract_document_types_from_query(self, query: str) -> list[str]:
        """Extract document type keywords from query."""
        doc_types = [
            '证据', '合同', '协议', '律师函', '起诉状', '起诉书',
            '欠条', '收据', '发票', '身份证', '判决书', '裁定书',
            '委托书', '授权书', '承诺书', '目录', '清单', '报告',
            'docx', 'pdf', 'jpg', 'png', '图片', '文件',
        ]
        found = []
        q = query.lower()
        for dt in doc_types:
            if dt in q:
                found.append(dt)
        return found

    def search_structured(self, query: str, k: int = 5) -> list[dict]:
        """Return search results as structured list instead of formatted string.
        
        Enhanced with:
        - Fuzzy matching with typo tolerance
        - Name + document type compound queries (e.g. "徐克林 证据")
        - Filename-based matching for document type queries
        - Multi-token partial matching
        - Better Chinese tokenization and relevance scoring
        """
        results = []
        seen = set()
        
        query_tokens = _segment(query)
        names = self._extract_names_from_query(query)
        doc_types = self._extract_document_types_from_query(query)
        
        # Boost exact phrase matches
        exact_matches = []
        if self._bm25 and self._tokenized_corpus:
            tokenized_query = query_tokens
            bm25_scores = self._bm25.get_scores(tokenized_query)
            
            # Build chunk index for BM25 results
            chunk_idx = 0
            for doc in self.documents:
                fn = doc.get("filename", "")
                for chunk in doc.get("chunks", []):
                    if chunk_idx < len(bm25_scores):
                        score = bm25_scores[chunk_idx]
                        if score > 0 or (len(self._tokenized_corpus) <= 3 and score > -5):
                            sig = (fn, chunk[:120])
                            if sig not in seen:
                                seen.add(sig)
                                
                                # Boost for exact phrase matches
                                boost = 1.0
                                if query.lower() in chunk.lower():
                                    boost = 2.0
                                elif any(token.lower() in chunk.lower() for token in query_tokens if len(token) > 1):
                                    boost = 1.5
                                
                                results.append({
                                    "filename": fn,
                                    "chunk": chunk,
                                    "score": round(float(score) * boost, 4),
                                    "match_mode": "bm25",
                                })
                    chunk_idx += 1
        
        # Filename matches with higher priority
        filename_matches = []
        for document in self.documents:
            fn = document.get("filename", "")
            fn_lower = fn.lower()
            fn_score = 0.0
            
            name_match = any(n in fn for n in names) if names else False
            doc_type_match = any(dt in fn_lower for dt in doc_types) if doc_types else False
            
            if name_match and doc_type_match:
                fn_score = 3.0
            elif name_match:
                fn_score = 2.0
            elif doc_type_match:
                fn_score = 1.5
            
            if fn_score > 0:
                for chunk in document.get("chunks", []):
                    sig = (fn, chunk[:120])
                    if sig not in seen:
                        seen.add(sig)
                        filename_matches.append({
                            "filename": fn,
                            "chunk": chunk,
                            "score": round(fn_score, 4),
                            "match_mode": "filename_match",
                        })
        
        results.extend(filename_matches)
        
        # Fuzzy matching fallback
        if not results:
            for document in self.documents:
                fn = document.get("filename")
                for chunk in document.get("chunks", []):
                    fuzzy = self._fuzzy_match(query, chunk)
                    if fuzzy > 0.3:
                        sig = (fn, chunk[:120])
                        if sig not in seen:
                            seen.add(sig)
                            results.append({
                                "filename": fn,
                                "chunk": chunk,
                                "score": round(fuzzy, 4),
                                "match_mode": "fuzzy",
                            })
        
        # Term frequency fallback
        if not results:
            for document in self.documents:
                fn = document.get("filename")
                for chunk in document.get("chunks", []):
                    score = self._tf_score(query_tokens, chunk)
                    if score > 0:
                        sig = (fn, chunk[:120])
                        if sig not in seen:
                            seen.add(sig)
                            results.append({
                                "filename": fn,
                                "chunk": chunk,
                                "score": round(float(score), 4),
                                "match_mode": "tf",
                            })
        
        # Sort by score descending
        results.sort(key=lambda item: item["score"], reverse=True)
        
        # Fill remaining with substring matches if needed
        if len(results) < k:
            skip = {(r["filename"], r["chunk"][:120]) for r in results}
            sub = self._substring_chunk_hits(query, k - len(results))
            for item in sub:
                sig = (item["filename"], item["chunk"][:120])
                if sig in skip:
                    continue
                skip.add(sig)
                results.append(item)
                if len(results) >= k:
                    break
        
        return results[:k]

    @staticmethod
    def _tf_score(query_tokens, chunk_text):
        if not query_tokens:
            return 0
        chunk_tokens = _segment(chunk_text)
        token_counts = Counter(chunk_tokens)
        return sum(token_counts.get(t, 0) for t in query_tokens)

    def purge_placeholder_documents(self) -> dict:
        """Remove RAG rows whose indexed text is OCR/read failure placeholders."""
        removed: list[str] = []
        kept: list = []
        for doc in self.documents:
            chunks = doc.get("chunks") or []
            combined = "\n".join(chunks)
            if _is_unindexable_placeholder_text(combined):
                fn = doc.get("filename") or ""
                if fn:
                    removed.append(fn)
                continue
            kept.append(doc)
        self.documents = kept
        self._save_store()
        self._rebuild_index()
        return {
            "removed": removed,
            "removed_count": len(removed),
            "remaining_count": len(kept),
        }

    def get_summary(self):
        total_chunks = sum(int(doc.get("chunk_count", 0) or 0) for doc in self.documents)
        return {
            "backend_mode": self.backend_mode,
            "ready": True,
            "error": self.init_error,
            "document_count": len(self.documents),
            "total_chunks": total_chunks,
            "documents": [
                {
                    "filename": doc.get("filename"),
                    "chunk_count": doc.get("chunk_count", 0),
                    "preview": doc.get("preview", ""),
                    "file_type": doc.get("file_type", "txt"),
                    "file_size_bytes": doc.get("file_size_bytes", 0),
                    "page_count": (doc.get("pdf_metadata") or {}).get("page_count"),
                }
                for doc in self.documents
            ],
        }


if __name__ == "__main__":
    rag = RAGManager()
    print(f"[RAG] Ready in {rag.backend_mode} mode. {rag.get_summary().get('total_chunks', 0)} chunks indexed.")
