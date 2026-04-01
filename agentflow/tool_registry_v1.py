import asyncio
import hashlib
import os
import time
from dataclasses import dataclass, field
from datetime import datetime, UTC
from enum import Enum
from typing import Any, Callable, Coroutine


class ToolMode(Enum):
    MANUAL = "manual"
    AUTO = "auto"


@dataclass
class ToolResult:
    tool_name: str
    success: bool
    latency_ms: float
    output: Any = None
    error: str | None = None
    mode_used: ToolMode = ToolMode.MANUAL


@dataclass
class ToolDefinition:
    name: str
    handler: Callable[..., Coroutine]
    mode: ToolMode = ToolMode.MANUAL
    description: str = ""
    requires_confirmation: bool = True
    call_count: int = 0
    success_count: int = 0
    total_latency_ms: float = 0.0


class ToolRegistryV1:
    def __init__(self, usage_tracker=None, default_mode: ToolMode = ToolMode.MANUAL):
        self.tracker = usage_tracker
        self.default_mode = default_mode
        self._tools: dict[str, ToolDefinition] = {}
        self._pending_confirmations: dict[str, ToolResult | None] = {}
        self._confirmation_events: dict[str, asyncio.Event] = {}
        self._register_builtin_tools()

    def _register_builtin_tools(self):
        self.register(
            "send_wechat",
            self._builtin_send_wechat,
            mode=ToolMode.MANUAL,
            description="Send a message via WeChat connector",
        )
        self.register(
            "publish_douyin",
            self._builtin_publish_douyin,
            mode=ToolMode.MANUAL,
            description="Publish content to Douyin",
        )
        self.register(
            "generate_pdf",
            self._builtin_generate_pdf,
            mode=ToolMode.AUTO,
            description="Generate a PDF document",
        )
        self.register(
            "generate_docx",
            self._builtin_generate_docx,
            mode=ToolMode.AUTO,
            description="Generate an editable DOCX legal document",
        )

    def register(
        self,
        name: str,
        handler: Callable[..., Coroutine],
        mode: ToolMode | None = None,
        description: str = "",
        requires_confirmation: bool | None = None,
    ):
        self._tools[name] = ToolDefinition(
            name=name,
            handler=handler,
            mode=mode or self.default_mode,
            description=description,
            requires_confirmation=requires_confirmation if requires_confirmation is not None else (mode == ToolMode.MANUAL),
        )

    def set_tool_mode(self, name: str, mode: ToolMode):
        if name in self._tools:
            self._tools[name].mode = mode
            self._tools[name].requires_confirmation = mode == ToolMode.MANUAL

    def set_all_mode(self, mode: ToolMode):
        for tool in self._tools.values():
            tool.mode = mode
            tool.requires_confirmation = mode == ToolMode.MANUAL

    def list_tools(self) -> list[dict]:
        return [
            {
                "name": t.name,
                "mode": t.mode.value,
                "description": t.description,
                "requires_confirmation": t.requires_confirmation,
                "stats": {
                    "calls": t.call_count,
                    "successes": t.success_count,
                    "success_rate": round(t.success_count / t.call_count, 2) if t.call_count else 0,
                    "avg_latency_ms": round(t.total_latency_ms / t.call_count, 1) if t.call_count else 0,
                },
            }
            for t in self._tools.values()
        ]

    async def execute(self, tool_name: str, **kwargs) -> ToolResult:
        tool = self._tools.get(tool_name)
        if tool is None:
            return ToolResult(
                tool_name=tool_name,
                success=False,
                latency_ms=0,
                error=f"Unknown tool: {tool_name}",
            )

        start = time.perf_counter()
        try:
            output = await tool.handler(**kwargs)
            latency = (time.perf_counter() - start) * 1000

            tool.call_count += 1
            tool.success_count += 1
            tool.total_latency_ms += latency

            result = ToolResult(
                tool_name=tool_name,
                success=True,
                latency_ms=latency,
                output=output,
                mode_used=tool.mode,
            )
        except Exception as e:
            latency = (time.perf_counter() - start) * 1000
            tool.call_count += 1
            tool.total_latency_ms += latency
            result = ToolResult(
                tool_name=tool_name,
                success=False,
                latency_ms=latency,
                error=str(e),
                mode_used=tool.mode,
            )

        if self.tracker:
            self.tracker.record_tool_call(
                tool_name=tool_name,
                success=result.success,
                latency_ms=result.latency_ms,
                extra={"mode": result.mode_used.value},
            )

        return result

    async def execute_with_confirmation(self, tool_name: str, **kwargs) -> ToolResult:
        tool = self._tools.get(tool_name)
        if tool is None:
            return await self.execute(tool_name, **kwargs)

        if not tool.requires_confirmation or tool.mode == ToolMode.AUTO:
            return await self.execute(tool_name, **kwargs)

        call_id = hashlib.md5(f"{tool_name}:{time.time()}".encode()).hexdigest()[:12]
        event = asyncio.Event()
        self._confirmation_events[call_id] = event
        self._pending_confirmations[call_id] = None

        pending_result = ToolResult(
            tool_name=tool_name,
            success=False,
            latency_ms=0,
            mode_used=ToolMode.MANUAL,
            error="awaiting_confirmation",
        )
        pending_result._call_id = call_id
        pending_result._kwargs = kwargs
        return pending_result

    async def confirm_and_execute(self, call_id: str) -> ToolResult | None:
        event = self._confirmation_events.pop(call_id, None)
        if event is None:
            return None
        pending = self._pending_confirmations.pop(call_id, None)
        if pending is None:
            return None

        tool_name = pending.tool_name
        kwargs = getattr(pending, "_kwargs", {})
        result = await self.execute(tool_name, **kwargs)
        event.set()
        return result

    async def reject(self, call_id: str) -> bool:
        event = self._confirmation_events.pop(call_id, None)
        self._pending_confirmations.pop(call_id, None)
        if event:
            event.set()
        return event is not None

    def get_pending(self) -> list[dict]:
        results = []
        for call_id, pending in self._pending_confirmations.items():
            if pending is None:
                continue
            results.append({
                "call_id": call_id,
                "tool_name": pending.tool_name,
                "kwargs": getattr(pending, "_kwargs", {}),
            })
        return results

    async def _builtin_send_wechat(self, **kwargs) -> dict:
        if os.getenv("AGENTFLOW_ENABLE_WECHAT", "0") != "1":
            return {"ok": False, "platform": "wechat", "error": "WeChat is paused (AGENTFLOW_ENABLE_WECHAT=0)"}
        await asyncio.sleep(0.01)
        return {
            "ok": True,
            "platform": "wechat",
            "contact_name_hash": hashlib.sha256(kwargs.get("contact_name", "").encode()).hexdigest()[:16],
        }

    async def _builtin_publish_douyin(self, **kwargs) -> dict:
        await asyncio.sleep(0.01)
        return {"ok": True, "platform": "douyin"}

    async def _builtin_generate_pdf(self, **kwargs) -> dict:
        return await asyncio.to_thread(self._generate_pdf_sync, **kwargs)

    def _generate_pdf_sync(self, **kwargs) -> dict:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            PageBreak, HRFlowable,
        )
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        title = kwargs.get("title", "Case Report")
        case_id = kwargs.get("case_id", "UNKNOWN")
        sections = kwargs.get("sections", [])
        highlights = kwargs.get("highlights", [])
        client_name = kwargs.get("client_name", "")
        now = datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC")

        report_dir = os.path.join(os.path.dirname(__file__), "data", "reports")
        os.makedirs(report_dir, exist_ok=True)
        safe_id = case_id.replace("/", "_")
        pdf_path = os.path.join(report_dir, f"{safe_id}_report.pdf")

        # Register a CJK-compatible font
        font_name = "Helvetica"
        cjk_font_path = None
        for candidate in [
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/System/Library/Fonts/Hiragino Sans GB.ttc",
        ]:
            if os.path.exists(candidate):
                cjk_font_path = candidate
                break

        doc = SimpleDocTemplate(
            pdf_path,
            pagesize=A4,
            leftMargin=25 * mm,
            rightMargin=25 * mm,
            topMargin=20 * mm,
            bottomMargin=20 * mm,
        )

        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            "CoverTitle",
            parent=styles["Title"],
            fontSize=22,
            spaceAfter=10,
            textColor=HexColor("#1e3a5f"),
        ))
        styles.add(ParagraphStyle(
            "CoverSub",
            parent=styles["Normal"],
            fontSize=11,
            textColor=HexColor("#64748b"),
            spaceAfter=4,
        ))
        styles.add(ParagraphStyle(
            "SectionHead",
            parent=styles["Heading2"],
            fontSize=14,
            textColor=HexColor("#2563eb"),
            spaceBefore=16,
            spaceAfter=8,
        ))
        styles.add(ParagraphStyle(
            "BodyText2",
            parent=styles["Normal"],
            fontSize=10,
            leading=14,
            spaceAfter=6,
        ))
        styles.add(ParagraphStyle(
            "HighlightBox",
            parent=styles["Normal"],
            fontSize=9,
            leading=13,
            leftIndent=10,
            rightIndent=10,
            spaceBefore=3,
            spaceAfter=3,
        ))
        styles.add(ParagraphStyle(
            "FooterStyle",
            parent=styles["Normal"],
            fontSize=8,
            textColor=HexColor("#94a3b8"),
            alignment=TA_RIGHT,
        ))

        CATEGORY_COLORS = {
            "dispute_clause": HexColor("#dc2626"),  # red
            "risk": HexColor("#ea580c"),             # orange
            "obligation": HexColor("#d97706"),       # amber
            "key_fact": HexColor("#2563eb"),          # blue
            "legal_basis": HexColor("#059669"),       # green
        }

        story = []

        # Cover page
        story.append(Spacer(1, 60 * mm))
        story.append(Paragraph(title, styles["CoverTitle"]))
        story.append(Spacer(1, 8 * mm))
        story.append(Paragraph(f"Case: {case_id}", styles["CoverSub"]))
        if client_name:
            story.append(Paragraph(f"Client: {client_name}", styles["CoverSub"]))
        story.append(Paragraph(f"Generated: {now}", styles["CoverSub"]))
        story.append(Spacer(1, 15 * mm))
        story.append(HRFlowable(width="60%", color=HexColor("#2563eb")))
        story.append(Spacer(1, 5 * mm))
        story.append(Paragraph("LegalAgent CRM - Automated Case Analysis", styles["CoverSub"]))
        story.append(PageBreak())

        # Sections
        for section in sections:
            heading = section.get("heading", "")
            body = section.get("body", "")
            if heading:
                story.append(Paragraph(heading, styles["SectionHead"]))
            if body:
                # Escape XML special chars for reportlab
                safe_body = body.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                # Split into paragraphs on double newlines
                for para in safe_body.split("\n\n"):
                    para = para.strip()
                    if para:
                        story.append(Paragraph(para, styles["BodyText2"]))

        # Highlights section
        if highlights:
            story.append(PageBreak())
            story.append(Paragraph("Document Highlights", styles["SectionHead"]))
            story.append(Paragraph(
                "Key passages identified by the AI for human review. "
                "Color coding: "
                '<font color="#dc2626">Red=Dispute Clause</font>, '
                '<font color="#d97706">Amber=Obligation</font>, '
                '<font color="#2563eb">Blue=Key Fact</font>, '
                '<font color="#059669">Green=Legal Basis</font>, '
                '<font color="#ea580c">Orange=Risk</font>.',
                styles["BodyText2"],
            ))
            story.append(Spacer(1, 4 * mm))

            for hl in highlights:
                text = hl.get("text", "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                category = hl.get("category", "key_fact")
                importance = hl.get("importance", "medium")
                reason = hl.get("reason", "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                color = CATEGORY_COLORS.get(category, HexColor("#64748b"))

                # Build highlight callout box
                importance_tag = f'<font color="#dc2626"><b>[{importance.upper()}]</b></font>' if importance == "high" else f'<font color="#94a3b8">[{importance}]</font>'
                category_tag = f'<font color="{color.hexval()}"><b>{category.replace("_", " ").title()}</b></font>'

                box_content = f'{importance_tag} {category_tag}<br/><br/><i>"{text}"</i>'
                if reason:
                    box_content += f'<br/><br/><font color="#94a3b8">Reason: {reason}</font>'

                box_table = Table(
                    [[Paragraph(box_content, styles["HighlightBox"])]],
                    colWidths=[doc.width],
                )
                box_table.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, -1), HexColor("#f8fafc")),
                    ("BOX", (0, 0), (-1, -1), 1.5, color),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                    ("TOPPADDING", (0, 0), (-1, -1), 8),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ]))
                story.append(box_table)
                story.append(Spacer(1, 3 * mm))

        # Footer with page numbers
        def add_footer(canvas, doc):
            canvas.saveState()
            canvas.setFont("Helvetica", 8)
            canvas.setFillColor(HexColor("#94a3b8"))
            canvas.drawRightString(
                A4[0] - 25 * mm,
                12 * mm,
                f"LegalAgent CRM | {case_id} | Page {canvas.getPageNumber()}",
            )
            canvas.drawString(
                25 * mm,
                12 * mm,
                now,
            )
            canvas.restoreState()

        doc.build(story, onFirstPage=add_footer, onLaterPages=add_footer)

        pages = doc.page
        return {
            "ok": True,
            "format": "pdf",
            "path": pdf_path,
            "filename": os.path.basename(pdf_path),
            "pages": pages,
        }

    async def _builtin_generate_docx(self, **kwargs) -> dict:
        return await asyncio.to_thread(self._generate_docx_sync, **kwargs)

    def _generate_docx_sync(self, **kwargs) -> dict:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        title = kwargs.get("title", "Case Report")
        case_id = kwargs.get("case_id", "UNKNOWN")
        sections = kwargs.get("sections", [])
        highlights = kwargs.get("highlights", [])
        client_name = kwargs.get("client_name", "")
        now = datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC")

        report_dir = os.path.join(os.path.dirname(__file__), "data", "reports")
        os.makedirs(report_dir, exist_ok=True)
        safe_id = case_id.replace("/", "_")
        docx_path = os.path.join(report_dir, f"{safe_id}_report.docx")

        doc = Document()

        doc.styles["Normal"].font.name = "PingFang SC"
        doc.styles["Normal"].font.size = Pt(11)
        doc.styles["Normal"].element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")

        section = doc.sections[0]
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

        # Cover page
        for _ in range(6):
            doc.add_paragraph("")
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(title)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

        doc.add_paragraph("")
        info_p = doc.add_paragraph()
        info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for line in [
            f"Case: {case_id}",
            f"Client: {client_name}" if client_name else "",
            f"Generated: {now}",
        ]:
            if line:
                run = info_p.add_run(line + "\n")
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        doc.add_paragraph("")
        footer_p = doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run("LegalAgent CRM — Automated Case Analysis")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        # Sections
        for section in sections:
            heading = section.get("heading", "")
            body = section.get("body", "")
            if heading:
                h = doc.add_heading(heading, level=2)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            if body:
                for para_text in body.split("\n\n"):
                    para_text = para_text.strip()
                    if para_text:
                        p = doc.add_paragraph(para_text)
                        for run in p.runs:
                            run.font.size = Pt(11)

        # Highlights
        if highlights:
            doc.add_page_break()
            h = doc.add_heading("Document Highlights", level=2)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            legend = doc.add_paragraph(
                "Key passages identified by the AI for human review. "
                "Color coding: Red=Dispute Clause, Amber=Obligation, "
                "Blue=Key Fact, Green=Legal Basis, Orange=Risk."
            )
            legend.runs[0].font.size = Pt(9)
            legend.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

            CATEGORY_COLORS = {
                "dispute_clause": RGBColor(0xDC, 0x26, 0x26),
                "risk": RGBColor(0xEA, 0x58, 0x0C),
                "obligation": RGBColor(0xD9, 0x77, 0x06),
                "key_fact": RGBColor(0x25, 0x63, 0xEB),
                "legal_basis": RGBColor(0x05, 0x96, 0x69),
            }
            for hl in highlights:
                text = hl.get("text", "")
                category = hl.get("category", "key_fact")
                importance = hl.get("importance", "medium")
                reason = hl.get("reason", "")
                color = CATEGORY_COLORS.get(category, RGBColor(0x64, 0x74, 0x8B))

                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.right_indent = Cm(0.5)

                imp_run = p.add_run(f"[{importance.upper()}] ")
                imp_run.bold = True
                imp_run.font.size = Pt(9)
                imp_run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26) if importance == "high" else RGBColor(0x94, 0xA3, 0xB8)

                cat_run = p.add_run(f"{category.replace('_', ' ').title()} — ")
                cat_run.bold = True
                cat_run.font.size = Pt(9)
                cat_run.font.color.rgb = color

                text_run = p.add_run(f'"{text}"')
                text_run.font.size = Pt(9)
                text_run.italic = True

                if reason:
                    r = p.add_run(f"\nReason: {reason}")
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

        # Page number footer
        section_el = doc.sections[0]
        footer = section_el.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = fp.add_run(f"LegalAgent CRM | {case_id} | Page ")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

        doc.save(docx_path)
        return {
            "ok": True,
            "format": "docx",
            "path": docx_path,
            "filename": os.path.basename(docx_path),
        }
